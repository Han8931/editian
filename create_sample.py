#!/usr/bin/env python3
"""Generate a sample .docx file with paragraphs, a table, and a chart."""

import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False
    print("Note: matplotlib not installed — chart will be skipped. Run: pip install matplotlib")


def add_chart(doc: Document) -> None:
    if not HAS_MATPLOTLIB:
        doc.add_paragraph("[Chart placeholder — install matplotlib to generate]")
        return

    quarters = ['Q1 2025', 'Q2 2025', 'Q3 2025', 'Q4 2025', 'Q1 2026']
    revenue  = [102, 119, 127, 141, 120]
    expenses = [84,   95, 101, 112,  91]

    x = list(range(len(quarters)))
    w = 0.35

    fig, ax = plt.subplots(figsize=(6.5, 3.8))
    bars_r = ax.bar([i - w / 2 for i in x], revenue,  w, label='Revenue',  color='#4472C4')
    bars_e = ax.bar([i + w / 2 for i in x], expenses, w, label='Expenses', color='#ED7D31')

    for bar in bars_r:
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1.5,
                f'${bar.get_height()}K', ha='center', va='bottom', fontsize=7.5)
    for bar in bars_e:
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1.5,
                f'${bar.get_height()}K', ha='center', va='bottom', fontsize=7.5)

    ax.set_title('Revenue vs. Expenses ($ thousands)', fontsize=11, pad=10)
    ax.set_xticks(x)
    ax.set_xticklabels(quarters, fontsize=9)
    ax.set_ylabel('Amount ($K)')
    ax.legend(fontsize=9)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(0, max(revenue) * 1.2)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)

    doc.add_picture(buf, width=Inches(5.8))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading('Q1 2026 Business Performance Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Executive Summary ──────────────────────────────────────────────────
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This report presents a comprehensive overview of the company's financial "
        "performance for the first quarter of 2026. Overall, the business demonstrated "
        "strong growth momentum, with total revenue increasing by 18% compared to the "
        "same period last year. Key drivers include the successful launch of our new "
        "product line and expanded market penetration in the Asia-Pacific region."
    )
    doc.add_paragraph(
        'Operating expenses were well-managed, with the cost-to-revenue ratio improving '
        'from 82% to 76% year-over-year. The company remains on track to meet its '
        'full-year targets, and the outlook for the remaining quarters is positive.'
    )

    # ── Financial Table ────────────────────────────────────────────────────
    doc.add_heading('Financial Overview', level=1)
    doc.add_paragraph(
        'The following table summarizes quarterly performance across key financial metrics:'
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ['Quarter', 'Revenue ($K)', 'Expenses ($K)', 'Net Profit ($K)', 'Margin (%)']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    rows = [
        ('Q1 2025', '102', '84',  '18', '17.6%'),
        ('Q2 2025', '119', '95',  '24', '20.2%'),
        ('Q3 2025', '127', '101', '26', '20.5%'),
        ('Q4 2025', '141', '112', '29', '20.6%'),
        ('Q1 2026', '120', '91',  '29', '24.2%'),
    ]
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()

    # ── Chart ──────────────────────────────────────────────────────────────
    doc.add_heading('Revenue vs. Expenses Chart', level=1)
    add_chart(doc)
    doc.add_paragraph()

    # ── Key Findings ───────────────────────────────────────────────────────
    doc.add_heading('Key Findings', level=1)
    findings = [
        'Revenue grew 17.6% year-over-year to $120K in Q1 2026, driven by strong demand in core product segments.',
        'Operating expenses decreased as a percentage of revenue, reflecting improved operational efficiency.',
        'Net profit margin reached 24.2%, the highest quarterly figure in company history.',
        'Customer acquisition cost dropped by 12% due to optimised digital marketing campaigns.',
        'Recurring subscription revenue now accounts for 62% of total revenue, up from 48% one year ago.',
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    # ── Recommendations ────────────────────────────────────────────────────
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        'Based on Q1 2026 performance, the leadership team recommends the following '
        'strategic priorities for the remainder of the fiscal year:'
    )
    recommendations = [
        'Accelerate hiring in the sales and engineering divisions to sustain growth trajectory.',
        'Increase R&D investment by 15% to maintain competitive advantage in core product areas.',
        'Expand the Asia-Pacific sales team to capture the growing demand identified in Q1.',
        'Initiate a cost review of legacy infrastructure to further improve the cost-to-revenue ratio.',
    ]
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {rec}')

    # ── Conclusion ─────────────────────────────────────────────────────────
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        "Q1 2026 results demonstrate the company's ability to grow revenue while "
        "simultaneously improving profitability. The trends observed this quarter "
        "are encouraging, and the management team is confident in achieving the "
        "full-year targets. Continued focus on operational excellence, customer "
        "retention, and strategic market expansion will be key to sustaining this "
        "momentum throughout FY 2026."
    )

    out = 'sample_report.docx'
    doc.save(out)
    print(f'Saved → {out}')


if __name__ == '__main__':
    main()
