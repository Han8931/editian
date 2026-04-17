import shutil
from typing import Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.run import Run


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_docx_revisions(
    input_path: str,
    output_path: str,
    revisions: list[dict[str, Any]],
) -> None:
    if input_path != output_path:
        shutil.copy2(input_path, output_path)

    doc = Document(output_path)
    structural_ops: list[dict[str, int]] = []

    def _current_index_for_original(original_index: int) -> int | None:
        current_index = original_index
        for op in structural_ops:
            if op["type"] == "delete_paragraph":
                deleted_index = op["index"]
                if deleted_index == original_index:
                    return None
                if deleted_index < original_index:
                    current_index -= 1
            elif op["type"] == "insert_paragraph":
                after_index = op["after_index"]
                if after_index < original_index:
                    current_index += 1
        return current_index

    def _current_insertion_anchor(after_original_index: int) -> int:
        if after_original_index < 0:
            return len(doc.paragraphs) - 1

        current_index = after_original_index
        for op in structural_ops:
            if op["type"] == "delete_paragraph":
                if op["index"] <= after_original_index:
                    current_index -= 1
            elif op["type"] == "insert_paragraph":
                if op["after_index"] <= after_original_index:
                    current_index += 1
        return current_index

    for revision in revisions:
        scope = revision["scope"]
        revised_text: str = revision["revised"]

        font_name: str | None  = revision.get("font_name")
        font_size: float | None = revision.get("font_size")
        align: str | None      = revision.get("align")
        bold: bool | None      = revision.get("bold")
        italic: bool | None    = revision.get("italic")
        underline: bool | None = revision.get("underline")
        strike: bool | None    = revision.get("strike")
        bullet: bool | None    = revision.get("bullet")

        if scope["type"] == "document":
            revised_paras = revised_text.split("\n\n")
            non_empty = [p for p in doc.paragraphs if p.text.strip()]
            for i, para in enumerate(non_empty):
                new_text = revised_paras[i] if i < len(revised_paras) else ""
                _set_paragraph_text(para, new_text, font_name, font_size, align, bold, italic, underline, strike, bullet)

        elif scope["type"] == "paragraphs":
            for idx in (scope.get("paragraph_indices") or []):
                current_idx = _current_index_for_original(idx)
                if current_idx is not None and 0 <= current_idx < len(doc.paragraphs):
                    _set_paragraph_text(doc.paragraphs[current_idx], revised_text, font_name, font_size, align, bold, italic, underline, strike, bullet)

        elif scope["type"] == "insert_paragraph":
            para_index = scope.get("paragraph_index", -1)
            current_anchor = _current_insertion_anchor(para_index)
            _insert_paragraph(doc, current_anchor, revised_text, font_name, font_size, align, bold, italic, underline, strike, bullet)
            structural_ops.append({"type": "insert_paragraph", "after_index": para_index})

        elif scope["type"] == "delete_paragraph":
            delete_indices = sorted(scope.get("paragraph_indices") or [], reverse=True)
            for idx in delete_indices:
                current_idx = _current_index_for_original(idx)
                if current_idx is None or not (0 <= current_idx < len(doc.paragraphs)):
                    continue
                _delete_paragraph(doc.paragraphs[current_idx])
                structural_ops.append({"type": "delete_paragraph", "index": idx})

        elif scope["type"] == "merge_paragraphs":
            merge_indices = sorted(scope.get("paragraph_indices") or [])
            if not merge_indices:
                continue

            target_idx = _current_index_for_original(merge_indices[0])
            if target_idx is None or not (0 <= target_idx < len(doc.paragraphs)):
                continue

            _set_paragraph_text(doc.paragraphs[target_idx], revised_text, font_name, font_size, align, bold, italic, underline, strike, bullet)

            for idx in merge_indices[1:]:
                current_idx = _current_index_for_original(idx)
                if current_idx is None or not (0 <= current_idx < len(doc.paragraphs)):
                    continue
                _delete_paragraph(doc.paragraphs[current_idx])
                structural_ops.append({"type": "delete_paragraph", "index": idx})

        elif scope["type"] == "insert_table":
            rows = scope.get("rows") or 2
            cols = scope.get("cols") or 2
            para_index = scope.get("paragraph_index", -1)
            current_anchor = _current_insertion_anchor(para_index)
            table = _insert_table(doc, current_anchor, rows, cols)
            table_cells = scope.get("cells") or []
            for row_index, row in enumerate(table_cells):
                if row_index >= len(table.rows):
                    break
                for cell_index, cell_text in enumerate(row):
                    if cell_index >= len(table.rows[row_index].cells):
                        break
                    cell = table.rows[row_index].cells[cell_index]
                    if cell.paragraphs:
                        _set_paragraph_text(cell.paragraphs[0], str(cell_text))

        elif scope["type"] == "table_cell":
            table_index = scope.get("table_index", 0)
            row_index = scope.get("row_index", 0)
            cell_index = scope.get("cell_index", 0)
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                if row_index < len(table.rows):
                    row = table.rows[row_index]
                    if cell_index < len(row.cells):
                        cell = row.cells[cell_index]
                        if cell.paragraphs:
                            _set_paragraph_text(cell.paragraphs[0], revised_text, font_name, font_size, align, bold, italic, underline, strike, bullet)

    doc.save(output_path)


def _insert_table(doc: Any, para_index: int, rows: int, cols: int) -> Any:
    """Insert a rows×cols table after the paragraph at para_index (-1 = end of document)."""
    # add_table appends to the body; we'll move it to the right position
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass

    tbl = table._tbl
    body = doc.element.body

    # Remove from where add_table placed it (end of body)
    body.remove(tbl)

    # Find the reference element to insert after
    paras = doc.paragraphs
    if para_index >= 0 and para_index < len(paras):
        ref_el = paras[para_index]._element
        ref_el.addnext(tbl)
    else:
        # Append before the last sectPr if present, otherwise at end
        sect_pr = body.find(qn('w:sectPr'))
        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), tbl)
        else:
            body.append(tbl)
    return table


def _insert_paragraph(doc: Any, para_index: int, text: str, font_name: str | None = None, font_size: float | None = None,
                      align: str | None = None, bold: bool | None = None, italic: bool | None = None,
                      underline: bool | None = None, strike: bool | None = None, bullet: bool | None = None) -> None:
    """Insert a paragraph after para_index (-1 = end of document)."""
    para = doc.add_paragraph()
    _set_paragraph_text(para, text, font_name, font_size, align, bold, italic, underline, strike, bullet)

    p_el = para._element
    body = doc.element.body
    body.remove(p_el)

    paras = doc.paragraphs
    if para_index >= 0 and para_index < len(paras):
        paras[para_index]._element.addnext(p_el)
    else:
        sect_pr = body.find(qn('w:sectPr'))
        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), p_el)
        else:
            body.append(p_el)


def _delete_paragraph(para: Any) -> None:
    p_el = para._element
    parent = p_el.getparent()
    if parent is not None:
        parent.remove(p_el)


def _visible_runs(para: Any) -> list:
    """Return all Run objects in *para*, including those nested inside
    <w:hyperlink> and <w:ins> elements, but excluding runs inside <w:del>
    (deleted/tracked-removal text that should not be edited)."""
    result = []
    for r in para._p.iter(qn('w:r')):
        # Walk up to see if this run lives inside a <w:del> ancestor.
        ancestor = r.getparent()
        in_del = False
        while ancestor is not None and ancestor is not para._p:
            if ancestor.tag == qn('w:del'):
                in_del = True
                break
            ancestor = ancestor.getparent()
        if not in_del:
            result.append(Run(r, para))
    return result


def _set_paragraph_text(para: Any, text: str, font_name: str | None = None, font_size: float | None = None,
                        align: str | None = None, bold: bool | None = None, italic: bool | None = None,
                        underline: bool | None = None, strike: bool | None = None, bullet: bool | None = None) -> None:
    # Use _visible_runs so we also reach runs nested inside <w:hyperlink> and
    # <w:ins> elements — para.runs only exposes direct <w:r> children, which
    # means paragraphs whose text lives exclusively in a hyperlink or tracked
    # insertion would otherwise have their new text *appended* (doubled) instead
    # of replacing the original.
    runs = _visible_runs(para)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        runs = [para.add_run(text)]

    if align is not None and align in _ALIGN_MAP:
        para.alignment = _ALIGN_MAP[align]

    if bullet is not None:
        _set_paragraph_bullet(para, bullet)

    for run in runs:
        if font_name is not None:
            run.font.name = font_name
        if font_size is not None:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if underline is not None:
            run.font.underline = underline
        if strike is not None:
            run.font.strike = strike


def _set_paragraph_bullet(para: Any, enabled: bool) -> None:
    p_pr = para._element.get_or_add_pPr()
    num_pr = p_pr.find(qn('w:numPr'))

    if enabled:
        try:
            para.style = 'List Bullet'
        except Exception:
            pass
        if num_pr is None:
            num_pr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            num_id = OxmlElement('w:numId')
            num_id.set(qn('w:val'), '1')
            num_pr.append(ilvl)
            num_pr.append(num_id)
            p_pr.append(num_pr)
        return

    if num_pr is not None:
        p_pr.remove(num_pr)
    try:
        if para.style and para.style.name.startswith('List '):
            para.style = 'Normal'
    except Exception:
        pass
