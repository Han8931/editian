import html as _html
import re
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Run
from typing import Any, Iterator


# ─── Constants ────────────────────────────────────────────────────────────────

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: 'left',
    WD_ALIGN_PARAGRAPH.CENTER: 'center',
    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: 'justify',
}

_HEADING_TAGS = {
    'Heading 1': 'h1', 'Heading 2': 'h2', 'Heading 3': 'h3',
    'Heading 4': 'h4', 'Heading 5': 'h5', 'Heading 6': 'h6',
    'Title': 'h1', 'Subtitle': 'h2',
}

# WD_COLOR_INDEX values → CSS hex
_HIGHLIGHT_COLORS = {
    1: '#ffff00', 2: '#00ff00', 3: '#00ffff', 4: '#ff00ff',
    5: '#0000ff', 6: '#ff0000', 7: '#00008b', 8: '#008b8b',
    9: '#006400', 10: '#800000', 11: '#808000', 12: '#ff69b4',
    13: '#808080', 14: '#d3d3d3',
}


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _rgb_hex(color_obj) -> str | None:
    try:
        if color_obj is not None and color_obj.type is not None:
            return f'#{color_obj.rgb}'
    except Exception:
        pass
    return None


def _iter_runs(para) -> Iterator[Run]:
    """Yield Run objects from a paragraph, including runs inside hyperlinks."""
    for child in para._p:
        if child.tag == qn('w:r'):
            yield Run(child, para)
        elif child.tag == qn('w:hyperlink'):
            url = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            for r in child.findall(qn('w:r')):
                yield _HyperlinkRun(r, para, url)
        elif child.tag == qn('w:ins'):
            # Tracked insertions — include the text
            for r in child.findall(qn('w:r')):
                yield Run(r, para)


class _HyperlinkRun(Run):
    def __init__(self, r, para, rel_id):
        super().__init__(r, para)
        self._rel_id = rel_id


def _render_run(run) -> str:
    text = run.text
    if not text:
        return ''

    escaped = _html.escape(text).replace('\n', '<br>')

    styles: list[str] = []

    # Font size
    try:
        if run.font.size:
            styles.append(f'font-size: {run.font.size.pt:.1f}pt')
    except Exception:
        pass

    # Font family
    try:
        if run.font.name:
            styles.append(f"font-family: '{run.font.name}'")
    except Exception:
        pass

    # Text color
    color = _rgb_hex(run.font.color)
    if color and color.upper() != '#000000':
        styles.append(f'color: {color}')

    # Highlight
    try:
        hl = run.font.highlight_color
        if hl is not None:
            val = hl.value if hasattr(hl, 'value') else int(hl)
            hl_css = _HIGHLIGHT_COLORS.get(val)
            if hl_css:
                styles.append(f'background-color: {hl_css}')
    except Exception:
        pass

    # Superscript / subscript
    try:
        if run.font.superscript:
            styles.append('vertical-align: super; font-size: 0.75em')
        elif run.font.subscript:
            styles.append('vertical-align: sub; font-size: 0.75em')
    except Exception:
        pass

    result = f'<span style="{"; ".join(styles)}">{escaped}</span>' if styles else escaped

    if run.bold:
        result = f'<strong>{result}</strong>'
    if run.italic:
        result = f'<em>{result}</em>'
    if run.underline:
        result = f'<u>{result}</u>'
    try:
        if run.font.strike:
            result = f'<s>{result}</s>'
    except Exception:
        pass

    # Wrap hyperlink runs in <a>
    if isinstance(run, _HyperlinkRun) and run._rel_id:
        try:
            url = run._p._parent.part.rels[run._rel_id].target_ref
            result = f'<a href="{_html.escape(url)}" style="color: #2563eb; text-decoration: underline;">{result}</a>'
        except Exception:
            pass

    return result


def _para_styles(para) -> tuple[str, str]:
    """Returns (html_tag, inline_style_string) for the paragraph element."""
    style_name = para.style.name if para.style else 'Normal'
    tag = _HEADING_TAGS.get(style_name, 'p')

    styles: list[str] = []
    pf = para.paragraph_format

    # Alignment
    try:
        align = _ALIGN_MAP.get(pf.alignment)
        if align:
            styles.append(f'text-align: {align}')
    except Exception:
        pass

    # Style-level font size (useful when runs don't set their own)
    try:
        if para.style and para.style.font and para.style.font.size:
            styles.append(f'font-size: {para.style.font.size.pt:.1f}pt')
    except Exception:
        pass

    # Indentation
    try:
        if pf.left_indent and pf.left_indent.pt:
            styles.append(f'padding-left: {pf.left_indent.pt:.1f}pt')
        if pf.first_line_indent and pf.first_line_indent.pt > 0:
            styles.append(f'text-indent: {pf.first_line_indent.pt:.1f}pt')
    except Exception:
        pass

    # Spacing
    try:
        if pf.space_before and pf.space_before.pt:
            styles.append(f'margin-top: {pf.space_before.pt:.1f}pt')
        if pf.space_after and pf.space_after.pt:
            styles.append(f'margin-bottom: {pf.space_after.pt:.1f}pt')
    except Exception:
        pass

    # Line spacing
    try:
        ls = pf.line_spacing
        if ls is not None:
            if hasattr(ls, 'pt'):
                styles.append(f'line-height: {ls.pt:.1f}pt')
            elif isinstance(ls, (int, float)):
                styles.append(f'line-height: {ls}')
    except Exception:
        pass

    return tag, '; '.join(styles)


def _style_list_info(para) -> dict | None:
    try:
        style_name = (para.style.name if para.style else '').strip()
    except Exception:
        style_name = ''

    if style_name:
        bullet_match = re.match(r'^List Bullet(?:\s+(\d+))?$', style_name, re.IGNORECASE)
        number_match = re.match(r'^List Number(?:\s+(\d+))?$', style_name, re.IGNORECASE)
        match = bullet_match or number_match
        if match:
            level = max(int(match.group(1) or '1') - 1, 0)
            return {
                'num_id': -abs(hash(style_name)),
                'level': level,
                'is_ordered': bool(number_match),
            }

    return None


def _get_list_info(para, doc) -> dict | None:
    """Returns {num_id, level, is_ordered} if para is a list item, else None."""
    try:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            return _style_list_info(para)
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return _style_list_info(para)
        numId_el = numPr.find(qn('w:numId'))
        if numId_el is None:
            return _style_list_info(para)
        num_id = int(numId_el.get(qn('w:val'), 0))
        if num_id == 0:
            return _style_list_info(para)
        ilvl_el = numPr.find(qn('w:ilvl'))
        level = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0

        is_ordered = False
        try:
            numbering = doc.part.numbering_part.numbering_definitions._numbering
            num_el = numbering.find(f'.//{qn("w:num")}[@{qn("w:numId")}="{num_id}"]')
            if num_el is not None:
                abs_id_el = num_el.find(qn('w:abstractNumId'))
                if abs_id_el is not None:
                    abs_id = abs_id_el.get(qn('w:val'))
                    abs_el = numbering.find(
                        f'.//{qn("w:abstractNum")}[@{qn("w:abstractNumId")}="{abs_id}"]'
                    )
                    if abs_el is not None:
                        lvl_el = abs_el.find(f'.//{qn("w:lvl")}[@{qn("w:ilvl")}="{level}"]')
                        if lvl_el is not None:
                            numFmt_el = lvl_el.find(qn('w:numFmt'))
                            if numFmt_el is not None:
                                fmt = numFmt_el.get(qn('w:val'), '')
                                is_ordered = fmt not in ('bullet', 'none', '')
        except Exception:
            pass

        return {'num_id': num_id, 'level': level, 'is_ordered': is_ordered}
    except Exception:
        return _style_list_info(para)


def _render_paragraph(para, index: int | None = None) -> str:
    tag, inline_style = _para_styles(para)
    style_attr = f' style="{inline_style}"' if inline_style else ''
    data_attr = f' data-para-index="{index}"' if index is not None else ''
    inner = ''.join(_render_run(r) for r in _iter_runs(para)) or '&nbsp;'
    return f'<{tag}{style_attr}{data_attr}>{inner}</{tag}>'


def _render_table(table, table_idx: int) -> str:
    rows: list[str] = []
    for row_idx, row in enumerate(table.rows):
        cells: list[str] = []
        for cell_idx, cell in enumerate(row.cells):
            cell_inner = ''.join(_render_paragraph(p) for p in cell.paragraphs)
            cells.append(
                f'<td data-cell-ref="t{table_idx}r{row_idx}c{cell_idx}" '
                'style="border: 1px solid #d1d5db; padding: 6px 12px; vertical-align: top; cursor: pointer;">'
                f'{cell_inner}</td>'
            )
        rows.append(f'<tr>{"".join(cells)}</tr>')
    return (
        '<table style="width: 100%; border-collapse: collapse; margin-bottom: 16px;">'
        f'{"".join(rows)}</table>'
    )


# ─── Public API ───────────────────────────────────────────────────────────────

def get_table_cells(file_path: str, table_index: int) -> list[dict]:
    """Return [{row, col, text}] for every non-empty cell in a table."""
    doc = Document(file_path)
    cells: list[dict] = []
    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
                if text:
                    cells.append({'row': row_idx, 'col': cell_idx, 'text': text})
    return cells


def get_cell_text(file_path: str, table_index: int, row_index: int, cell_index: int) -> str | None:
    doc = Document(file_path)
    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        if row_index < len(table.rows):
            row = table.rows[row_index]
            if cell_index < len(row.cells):
                cell = row.cells[cell_index]
                return '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
    return None


def get_docx_page_count(file_path: str) -> int:
    """
    Return the page count for the DOCX file.
    Primary:  docProps/app.xml <Pages> element (set by Word on save).
    Fallback: count <w:lastRenderedPageBreak> elements in document.xml
              (page-break markers Word embeds in the body XML) + 1.
    """
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            names = z.namelist()

            # Primary: app.xml metadata
            if 'docProps/app.xml' in names:
                root = ET.fromstring(z.read('docProps/app.xml'))
                ns = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                el = root.find('ep:Pages', ns)
                if el is not None and el.text:
                    count = int(el.text)
                    if count > 0:
                        return count

            # Fallback: lastRenderedPageBreak markers in document.xml
            doc_entry = next((n for n in names if n.endswith('word/document.xml')), None)
            if doc_entry:
                xml = z.read(doc_entry).decode('utf-8', errors='replace')
                breaks = xml.count('<w:lastRenderedPageBreak')
                if breaks > 0:
                    return breaks + 1
    except Exception:
        pass
    return 1


def parse_docx(file_path: str) -> dict[str, Any]:
    doc = Document(file_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
    return {"paragraphs": paragraphs, "total": len(doc.paragraphs)}


def get_docx_html(file_path: str) -> str:
    doc = Document(file_path)

    # Map XML element → (paragraph object, index) for body-level paragraphs only
    para_map = {p._element: (p, i) for i, p in enumerate(doc.paragraphs)}
    # Map XML element → table object
    table_map = {t._tbl: t for t in doc.tables}

    parts: list[str] = []
    table_counter = 0
    # Stack entries: (html_tag, num_id, level)
    list_stack: list[tuple[str, int, int]] = []

    def close_lists_deeper_than(target_level: int) -> None:
        while list_stack and list_stack[-1][2] > target_level:
            parts.append(f'</{list_stack.pop()[0]}>')

    def close_lists_at_or_deeper_than(target_level: int) -> None:
        while list_stack and list_stack[-1][2] >= target_level:
            parts.append(f'</{list_stack.pop()[0]}>')

    def close_all_lists() -> None:
        while list_stack:
            parts.append(f'</{list_stack.pop()[0]}>')

    for child in doc.element.body:
        tag_name = child.tag

        if tag_name == qn('w:p'):
            entry = para_map.get(child)
            if entry is None:
                continue
            para, idx = entry

            list_info = _get_list_info(para, doc)

            if list_info:
                level = list_info['level']
                list_tag = 'ol' if list_info['is_ordered'] else 'ul'
                list_id = list_info['num_id']

                close_lists_deeper_than(level)

                if (
                    list_stack
                    and list_stack[-1][2] == level
                    and (list_stack[-1][0] != list_tag or list_stack[-1][1] != list_id)
                ):
                    close_lists_at_or_deeper_than(level)

                if not list_stack or list_stack[-1][2] < level:
                    list_stack.append((list_tag, list_id, level))
                    indent_px = (level + 1) * 24
                    parts.append(
                        f'<{list_tag} style="padding-left: {indent_px}px; margin: 4px 0;">'
                    )

                data_attr = f' data-para-index="{idx}"' if para.text.strip() else ''
                inner = ''.join(_render_run(r) for r in _iter_runs(para)) or '&nbsp;'
                parts.append(f'<li{data_attr}>{inner}</li>')

            else:
                close_all_lists()
                index = idx if para.text.strip() else None
                parts.append(_render_paragraph(para, index))

        elif tag_name == qn('w:tbl'):
            close_all_lists()
            table = table_map.get(child)
            if table is not None:
                parts.append(_render_table(table, table_counter))
                table_counter += 1

    close_all_lists()
    return ''.join(parts)
